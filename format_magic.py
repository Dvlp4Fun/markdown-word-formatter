"""
format_magic.py
Scans ./input for .docx files, lets the user choose one in the terminal,
then generates a "Format Magic" styled version into ./output.
Requires: pip install python-docx
"""

import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT_DIR = "input"
OUTPUT_DIR = "output"


# ---------------------------------------------------------------- helpers
def add_bottom_border(p, color="CCCCCC", size="6"):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr'); b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), size)
    b.set(qn('w:space'), '4'); b.set(qn('w:color'), color)
    pBdr.append(b); pPr.append(pBdr)


def shade(p, fill="F5F5F5"):
    pPr = p._p.get_or_add_pPr(); s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear'); s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), fill); pPr.append(s)


def add_runs_with_bold(paragraph, text, size=11):
    for part in re.split(r'(\*\*.*?\*\*)', text):
        if part.startswith('**') and part.endswith('**'):
            r = paragraph.add_run(part[2:-2]); r.bold = True
        else:
            r = paragraph.add_run(part)
        r.font.size = Pt(size)


# ---------------------------------------------------------------- builders
def add_heading(doc, text, level):
    sizes = {1: 22, 2: 16, 3: 13, 4: 11.5}
    p = doc.add_paragraph()
    r = p.add_run(text.strip()); r.bold = True
    r.font.size = Pt(sizes.get(level, 11))
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    if level <= 2:
        add_bottom_border(p, "CCCCCC" if level == 1 else "E0E0E0",
                          "6" if level == 1 else "4")


def add_body(doc, text):
    p = doc.add_paragraph()
    add_runs_with_bold(p, text)
    for r in p.runs:
        if not r.bold:
            r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_bullet(doc, text):
    add_runs_with_bold(doc.add_paragraph(style="List Bullet"), text)


def add_numbered(doc, text):
    add_runs_with_bold(doc.add_paragraph(style="List Number"), text)


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    r = p.add_run("\n".join(lines))
    r.font.name = "Consolas"; r.font.size = Pt(8)
    shade(p, "F5F5F5")


def add_pipe_table(doc, rows):
    cols = max(len(r) for r in rows)
    rows = [r + [""] * (cols - len(r)) for r in rows]
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            cell.paragraphs[0].text = ""
            add_runs_with_bold(cell.paragraphs[0], val, size=9)
            if i == 0:
                for rn in cell.paragraphs[0].runs:
                    rn.bold = True


# ---------------------------------------------------------------- detectors
H_RE   = re.compile(r'^(#{1,6})\s+(.*)$')
BUL_RE = re.compile(r'^\s*[-*]\s+(.*)$')
NUM_RE = re.compile(r'^\s*\d+\.\s+(.*)$')
FENCE  = re.compile(r'^\s*```')
HR_RE  = re.compile(r'^\s*---\s*$')
SEP_RE = re.compile(r'^\s*\|?[\s:|-]+\|?\s*$')
PIPE   = lambda s: s.strip().startswith('|') and s.strip().endswith('|')


def split_pipe(line):
    return [c.strip() for c in line.strip().strip('|').split('|')]


# ---------------------------------------------------------------- build
def build(inp, out):
    lines = [p.text for p in Document(inp).paragraphs]
    doc = Document(); n = len(lines); i = 0

    while i < n:
        line = lines[i].strip()
        if not line:
            i += 1; continue

        if FENCE.match(line):
            block = []; i += 1
            while i < n and not FENCE.match(lines[i]):
                block.append(lines[i]); i += 1
            i += 1
            add_code_block(doc, block); continue

        if HR_RE.match(line):
            i += 1; continue

        h = H_RE.match(line)
        if h:
            add_heading(doc, h.group(2), len(h.group(1))); i += 1; continue

        if PIPE(line):
            rows = []
            while i < n and PIPE(lines[i]):
                if not SEP_RE.match(lines[i]):
                    rows.append(split_pipe(lines[i]))
                i += 1
            if rows:
                add_pipe_table(doc, rows)
            continue

        m = NUM_RE.match(line)
        if m:
            add_numbered(doc, m.group(1)); i += 1; continue
        m = BUL_RE.match(line)
        if m:
            add_bullet(doc, m.group(1)); i += 1; continue

        add_body(doc, line); i += 1

    doc.save(out)


# ---------------------------------------------------------------- menu
def main():
    os.makedirs(INPUT_DIR, exist_ok=True)
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    files = sorted(f for f in os.listdir(INPUT_DIR) if f.lower().endswith(".docx"))
    if not files:
        print(f"No .docx files found in '{INPUT_DIR}/'. Add files and re-run.")
        return

    print("\n=== Format Magic - Document Formatter ===")
    print(f"Files found in '{INPUT_DIR}/':\n")
    for idx, f in enumerate(files, start=1):
        print(f"  [{idx}] {f}")
    print(f"  [0] Format ALL files")

    # get user choice
    while True:
        choice = input("\nEnter the number of the file to format: ").strip()
        if not choice.isdigit():
            print("Please enter a valid number.")
            continue
        choice = int(choice)
        if 0 <= choice <= len(files):
            break
        print(f"Please enter a number between 0 and {len(files)}.")

    # process
    targets = files if choice == 0 else [files[choice - 1]]
    for fname in targets:
        in_path = os.path.join(INPUT_DIR, fname)
        base = os.path.splitext(fname)[0]
        out_path = os.path.join(OUTPUT_DIR, f"{base}_formatted.docx")
        build(in_path, out_path)
        print(f"  ✔ {fname}  ->  {out_path}")

    print("\nDone! Check the 'output/' folder.\n")


if __name__ == "__main__":
    main()